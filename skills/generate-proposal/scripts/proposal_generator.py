#!/usr/bin/env python3
"""
Proposal Generator - Convert markdown proposals to professionally formatted Word documents.
Generates .docx files with styled headers, bullet points, and horizontal lines.
"""

import sys
import os
import argparse
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph in docx."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to docx with formatting."""
    if not HAS_DOCX:
        print(f"Skipping .docx generation: python-docx not installed.")
        print(f"Install with: pip install python-docx")
        return False

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue

        # Headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.size = Pt(24)
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            elif level == 2:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            elif level == 3:
                p = doc.add_heading(text, level=3)
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                p.runs[0].font.bold = True
            else:
                p = doc.add_paragraph(text)
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(12)

        # Bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)

        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                p.add_run(text)

        # Numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                p.add_run(text)

        # Regular paragraph
        else:
            p = doc.add_paragraph(line)

        i += 1

    doc.save(docx_file)
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Convert markdown proposals to Word documents'
    )
    parser.add_argument('input', help='Path to markdown proposal file')
    parser.add_argument('--output', '-o',
                        help='Output .docx path (default: same name as input)')

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix('.docx')

    print(f"Converting: {input_path} -> {output_path}")
    success = markdown_to_docx(str(input_path), str(output_path))

    if success:
        print(f"Saved: {output_path}")
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
